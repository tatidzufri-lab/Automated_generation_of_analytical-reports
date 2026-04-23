from __future__ import annotations

import logging
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from app.core.config import Settings, get_settings
from app.services.artifact_naming import humanize_artifact_name
from app.services.file_service import FileService, StoredFile
from app.services.time_utils import filename_timestamp, format_local


logger = logging.getLogger(__name__)


class ReportService:
    def __init__(self, file_service: FileService, settings: Settings | None = None) -> None:
        self.file_service = file_service
        self.settings = settings or get_settings()

    def generate_report(
        self,
        stored_file: StoredFile,
        analysis: dict[str, Any],
        charts: list[dict[str, Any]],
        preview_context: dict[str, Any],
        business_metrics: dict[str, Any] | None = None,
    ) -> dict[str, str]:
        self.file_service.ensure_storage()
        report_name = self._build_report_name(stored_file.file_id)
        report_path = self.settings.output_dir / report_name

        document = Document()
        title = document.add_heading("Analytics Assistant Report", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        document.add_paragraph(f"Исходный файл: {stored_file.original_name}")
        document.add_paragraph(f"Тип: {stored_file.kind}")
        document.add_paragraph(f"Сформирован: {format_local('%d.%m.%Y %H:%M')}")

        if business_metrics and (business_metrics.get("total_sales") or business_metrics.get("total_orders")):
            document.add_heading("Бизнес-метрики", level=1)
            business_table = document.add_table(rows=1, cols=2)
            business_table.style = "Light Grid Accent 3"
            header = business_table.rows[0].cells
            header[0].text = "Показатель"
            header[1].text = "Значение"
            for label, value in (
                ("Общая сумма", f"{business_metrics.get('total_sales', 0):,.2f}".replace(",", " ")),
                ("Средний чек", f"{business_metrics.get('avg_ticket', 0):,.2f}".replace(",", " ")),
                ("Количество записей", f"{business_metrics.get('total_orders', 0):,}".replace(",", " ")),
            ):
                row = business_table.add_row().cells
                row[0].text = label
                row[1].text = value

            top_items = business_metrics.get("top_items") or []
            if top_items:
                document.add_heading("Топ позиций", level=2)
                top_table = document.add_table(rows=1, cols=2)
                top_table.style = "Light Shading Accent 3"
                top_header = top_table.rows[0].cells
                top_header[0].text = "Позиция"
                top_header[1].text = "Сумма"
                for item in top_items:
                    row = top_table.add_row().cells
                    row[0].text = str(item.get("item", ""))
                    row[1].text = f"{float(item.get('amount', 0)):,.2f}".replace(",", " ")

        document.add_heading("Краткое резюме", level=1)
        summary_table = document.add_table(rows=1, cols=2)
        summary_table.style = "Light Grid Accent 1"
        header_cells = summary_table.rows[0].cells
        header_cells[0].text = "Метрика"
        header_cells[1].text = "Значение"
        for label, value in {
            "Rows / Height": analysis["summary"]["rows"],
            "Columns / Width": analysis["summary"]["columns"],
            "Numeric metrics": analysis["summary"]["numeric_columns"],
            "Missing cells": analysis["summary"]["missing_cells"],
        }.items():
            row = summary_table.add_row().cells
            row[0].text = str(label)
            row[1].text = str(value)

        document.add_heading("Выводы", level=1)
        for insight in analysis["insights"]:
            document.add_paragraph(insight, style="List Bullet")

        if analysis["stats"]:
            document.add_heading("Статистика", level=1)
            headers = list(analysis["stats"][0].keys())
            stats_table = document.add_table(rows=1, cols=len(headers))
            stats_table.style = "Light Shading Accent 1"
            for index, header in enumerate(headers):
                stats_table.rows[0].cells[index].text = str(header)
            for item in analysis["stats"]:
                cells = stats_table.add_row().cells
                for index, header in enumerate(headers):
                    cells[index].text = str(item[header])

        if analysis["missing_summary"]:
            document.add_heading("Пропуски", level=1)
            missing_table = document.add_table(rows=1, cols=3)
            missing_table.style = "Light Grid Accent 2"
            for index, header in enumerate(["Column", "Missing", "Percent"]):
                missing_table.rows[0].cells[index].text = header
            for item in analysis["missing_summary"]:
                cells = missing_table.add_row().cells
                cells[0].text = item["column"]
                cells[1].text = str(item["missing"])
                cells[2].text = item["percent"]

        document.add_heading("Preview", level=1)
        if stored_file.kind == "table":
            headers = preview_context["table_headers"][:6]
            preview_table = document.add_table(rows=1, cols=len(headers))
            preview_table.style = "Table Grid"
            for index, header in enumerate(headers):
                preview_table.rows[0].cells[index].text = str(header)
            for row_values in preview_context["table_rows"][:10]:
                cells = preview_table.add_row().cells
                for index, value in enumerate(row_values[: len(headers)]):
                    cells[index].text = str(value)
        else:
            document.add_paragraph(
                f"Изображение {preview_context['image_width']}×{preview_context['image_height']} ({preview_context['image_mode']})"
            )

        if charts:
            document.add_heading("Графики", level=1)
            for chart in charts:
                display_name = (
                    chart.get("display_name")
                    or humanize_artifact_name(chart.get("file_name", ""))
                )
                caption = document.add_paragraph()
                caption_run = caption.add_run(display_name)
                caption_run.bold = True
                description = chart.get("description")
                if description:
                    document.add_paragraph(description)
                document.add_picture(
                    str(self.settings.output_dir / chart["file_name"]),
                    width=Inches(6.3),
                )

        document.save(report_path)
        logger.info("Generated report %s for %s", report_name, stored_file.file_id)

        return {
            "file_name": report_name,
            "display_name": humanize_artifact_name(report_name),
            "storage_url": f"/storage/outputs/{report_name}",
            "download_url": f"/download/{report_name}",
        }

    def _build_report_name(self, file_id: str) -> str:
        return f"{file_id}__report__{filename_timestamp()}.docx"
